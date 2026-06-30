"""Extract plain text from issue attachments (PDF / Word / Excel / text).

Lets the AI read attachment content without spending tokens on raw binaries.
Pure-python parsers (pypdf, python-docx, openpyxl); image OCR is intentionally
out of scope for now (would need tesseract) — flagged via `extractable`.
"""

from __future__ import annotations

import csv
import io
import os
import time
from typing import Any

import structlog

log = structlog.get_logger(__name__)

# Cap extracted text so a huge spreadsheet/PDF can't blow up the LLM prompt.
# Многоактные сканы (Волжское ПО: ~1.1k символов на акт) при 8000 обрезались на
# ~6 актах — подняли, чтобы пакетный разбор видел ВСЕ акты в одном PDF.
_MAX_CHARS = 50000

_TEXT_EXTS = {".txt", ".csv", ".log"}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
_XLSX_EXTS = {".xlsx", ".xlsm"}
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".heic", ".heif", ".bmp", ".tif", ".tiff"}


def _ext(filename: str) -> str:
    i = filename.rfind(".")
    return filename[i:].lower() if i >= 0 else ""


# OCR (scanned PDFs / photos). Best-effort: needs tesseract + PyMuPDF.
# OCR — синхронная CPU-bound работа; вызывающий выносит её в поток (asyncio.to_thread),
# чтобы не блокировать единственный event loop uvicorn (иначе ВЕСЬ сервис висит и
# CPU одного ядра = 100% на время разбора).
# Потолок страниц + БЮДЖЕТ ВРЕМЕНИ гарантируют, что разбор многостраничного скана
# вернётся задолго до фронт-таймаута (90 с) даже на медленном сервере под свопом.
# Раньше 40 страниц @200 DPI могли занять >90 с → запрос отваливался, результат не
# кэшировался, оператор бился в цикл «Ошибка разбора». Лучше разобрать первые N актов,
# чем не разобрать ничего.
_OCR_LANG = "rus+eng"
# Сервер: 1 ядро, ~1 ГБ RAM. 150 DPI вместо 200 — вдвое меньше пикселей: ~4 с/стр.
# вместо ~8 с и заметно меньше пиковой памяти на пиксмап. Точность распознавания
# гос.номеров на печатных актах при 150 DPI сохраняется.
_OCR_DPI = 150
# Сколько страниц максимум OCR-им за ОДНО окно (один вызов). Ограничивает пиковую
# память и длительность; остаток документа дораспознаётся при следующем вызове
# (см. ocr_pdf_window — возобновляемый OCR с курсором страниц).
_OCR_WINDOW_PAGES = 12
# Бюджет времени на одно окно OCR. Возврат гарантированно укладывается в окно
# запроса (~90 с) с запасом на анализ телеметрии. Что не успели — доберётся при
# повторном «Обновить разбор» (накопительно, из кэша по вложению).
_OCR_TIME_BUDGET_S = 45.0


def _ocr_available() -> bool:
    import importlib.util as u
    return bool(u.find_spec("pytesseract") and u.find_spec("PIL"))


def _heif_available() -> bool:
    """pillow-heif нужен, чтобы PIL открывал HEIC/HEIF (фото с iPhone)."""
    import importlib.util as u
    return bool(u.find_spec("pillow_heif"))


def is_extractable(filename: str) -> bool:
    e = _ext(filename)
    if e in _TEXT_EXTS | _PDF_EXTS | _DOCX_EXTS | _XLSX_EXTS:
        return True
    # HEIC/HEIF — только если установлен pillow-heif (иначе PIL не откроет).
    if e in {".heic", ".heif"}:
        return _ocr_available() and _heif_available()
    # Прочие картинки — когда доступен OCR.
    return e in _IMAGE_EXTS and _ocr_available()


def kind(filename: str) -> str:
    e = _ext(filename)
    if e in _PDF_EXTS:
        return "pdf"
    if e in _DOCX_EXTS:
        return "word"
    if e in _XLSX_EXTS:
        return "excel"
    if e in _IMAGE_EXTS:
        return "image"
    if e in _TEXT_EXTS:
        return "text"
    return "other"


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) > _MAX_CHARS:
        return text[:_MAX_CHARS] + "\n…[текст обрезан]"
    return text


def extract_text(filename: str, data: bytes) -> str:
    """Return extracted text, or '' if not extractable / on parse error."""
    e = _ext(filename)
    try:
        if e in _PDF_EXTS:
            return _truncate(_pdf(data))
        if e in _DOCX_EXTS:
            return _truncate(_docx(data))
        if e in _XLSX_EXTS:
            return _truncate(_xlsx(data))
        if e in _TEXT_EXTS:
            return _truncate(_plain(data))
        if e in _IMAGE_EXTS:
            return _truncate(_ocr_image(data))
    except Exception:  # pragma: no cover - never break the caller
        log.warning("attachment_extract_failed", filename=filename)
    return ""


def _cyrillic_ratio(s: str) -> float:
    letters = [c for c in s if c.isalpha()]
    if not letters:
        return 0.0
    cyr = sum(1 for c in letters if "а" <= c.lower() <= "я" or c.lower() == "ё")
    return cyr / len(letters)


def _pdf_text_layer(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    return "\n".join(parts)


def _pdf(data: bytes) -> str:
    text = _pdf_text_layer(data)
    # Scanned PDFs (no text layer) come back empty; some have broken font
    # encodings that decode to Latin mojibake. Our domain is Russian docs —
    # if the text layer is missing/unreliable (almost no Cyrillic), fall back
    # to OCR rendering the pages as images (одно окно — без возобновления).
    if len(text) < 40 or _cyrillic_ratio(text) < 0.15:
        ocr, _next, _total = _pdf_ocr_window(data, 0)
        if ocr.strip():
            return ocr
        return "" if _cyrillic_ratio(text) < 0.15 else text
    return text


def _ensure_tesseract_cmd() -> None:
    """Point pytesseract at the tesseract binary explicitly — the systemd
    service runs with a minimal PATH (venv only), so auto-discovery fails."""
    import shutil

    import pytesseract

    current = getattr(pytesseract.pytesseract, "tesseract_cmd", "tesseract")
    if current and current != "tesseract" and os.path.exists(current):
        return
    cmd = shutil.which("tesseract") or "/usr/bin/tesseract"
    pytesseract.pytesseract.tesseract_cmd = cmd


def _ocr_image(data: bytes) -> str:
    if not _ocr_available():
        return ""
    import pytesseract
    from PIL import Image

    # Зарегистрировать HEIC/HEIF-декодер, если установлен pillow-heif (фото iPhone).
    try:
        import pillow_heif
        pillow_heif.register_heif_opener()
    except Exception:
        pass

    _ensure_tesseract_cmd()
    img = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(img, lang=_OCR_LANG)


def _pdf_ocr_window(data: bytes, start_page: int = 0) -> tuple[str, int, int]:
    """OCR a time/page-bounded WINDOW of a scanned PDF, starting at ``start_page``.

    Возобновляемый OCR: сервер слабый (1 ядро/1 ГБ), большой многоактный PDF не
    влезает в одно окно запроса. За вызов распознаём до ``_OCR_WINDOW_PAGES`` стр.
    или пока не истёк ``_OCR_TIME_BUDGET_S``, начиная с ``start_page``. Возвращаем
    ``(text_окна, next_page, total_pages)``. ``next_page >= total_pages`` означает,
    что документ распознан полностью; иначе вызывающий сохраняет курсор и дотягивает
    остаток следующим вызовом (накопление текста — на стороне caller).
    """
    import importlib.util as u
    if not _ocr_available() or not u.find_spec("fitz"):
        return "", start_page, start_page
    import fitz  # PyMuPDF

    parts: list[str] = []
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        total = doc.page_count
        zoom = _OCR_DPI / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        deadline = time.monotonic() + _OCR_TIME_BUDGET_S
        page = start_page
        done_in_window = 0
        while page < total and done_in_window < _OCR_WINDOW_PAGES:
            # Бюджет проверяем ПОСЛЕ первой страницы окна — хотя бы одну страницу
            # обрабатываем всегда, как бы ни был медленен сервер.
            if done_in_window and time.monotonic() > deadline:
                break
            pix = doc[page].get_pixmap(matrix=matrix)
            t = _ocr_image(pix.tobytes("png"))
            if t.strip():
                parts.append(t)
            page += 1
            done_in_window += 1
        if page < total:
            log.warning("pdf_ocr_window_partial", start=start_page,
                        next_page=page, total=total)
        return "\n".join(parts), page, total
    finally:
        doc.close()


def extract_resumable(filename: str, data: bytes,
                      start_page: int = 0) -> tuple[str, bool, int]:
    """Извлечь текст вложения с поддержкой возобновления для сканов PDF.

    Возвращает ``(text, complete, next_page)``:
    - не-PDF (docx/xlsx/text/image) и PDF с текстовым слоем — один проход,
      ``complete=True``, ``next_page=0`` (текст полный, обрезан до ``_MAX_CHARS``);
    - скан PDF — OCR ОКНА начиная с ``start_page``; ``text`` — только этот фрагмент
      (caller накапливает), ``complete`` истинно когда документ распознан целиком,
      ``next_page`` — курсор для следующего вызова.
    На ошибке/неизвлекаемом типе → ``("", True, start_page)``.
    """
    e = _ext(filename)
    try:
        if e in _PDF_EXTS:
            if start_page == 0:
                layer = _pdf_text_layer(data)
                if len(layer) >= 40 and _cyrillic_ratio(layer) >= 0.15:
                    return _truncate(layer), True, 0
            ocr, next_page, total = _pdf_ocr_window(data, start_page)
            return _truncate(ocr), next_page >= total, next_page
        if e in _DOCX_EXTS:
            return _truncate(_docx(data)), True, 0
        if e in _XLSX_EXTS:
            return _truncate(_xlsx(data)), True, 0
        if e in _TEXT_EXTS:
            return _truncate(_plain(data)), True, 0
        if e in _IMAGE_EXTS:
            return _truncate(_ocr_image(data)), True, 0
    except Exception:  # pragma: no cover - never break the caller
        log.warning("attachment_extract_failed", filename=filename)
    return "", True, start_page


def _docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Лист: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(c) for c in row if c is not None]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def _plain(data: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")
